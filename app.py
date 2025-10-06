# app.py
import streamlit as st
import pandas as pd
import numpy as np
import re
from io import BytesIO
from datetime import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ReportLab for PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import utils
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

# python-docx for editable docx
from docx import Document
from docx.shared import Inches

# ---------- CONFIG ----------
ACADEMY_NAME_DEFAULT = "Avenir Academy"
STRONG_TH = 70.0
AVERAGE_LO = 40.0
TOP_COUNT = 3
BOTTOM_COUNT = 3

IGNORE_COLS_DEFAULT = [
    "Student Id Fk", "Roll Number", "Section Roll Number",
    "Student Name", "Father Name", "Class Name", "Section Name",
    "Evaluation", "Total Marks", "Obtained Marks",
    "Current Percentage", "Current Position", "Current Grade"
]
# ----------------------------

st.set_page_config(page_title="Student Performance Analyzer", page_icon="📊", layout="wide")
st.title("📊 Student Performance Analyzer")

# Initialize session state for comparison mode
if 'comparison_mode' not in st.session_state:
    st.session_state.comparison_mode = False
if 'df_week1' not in st.session_state:
    st.session_state.df_week1 = None
if 'df_week2' not in st.session_state:
    st.session_state.df_week2 = None
if 'comparison_data' not in st.session_state:
    st.session_state.comparison_data = None

# Sidebar options
st.sidebar.header("Report Settings")
academy_name = st.sidebar.text_input("Academy/School name", value=ACADEMY_NAME_DEFAULT)

# Comparison mode toggle
comparison_mode = st.sidebar.checkbox("📈 Enable Two-Week Comparison", value=st.session_state.comparison_mode)
if comparison_mode != st.session_state.comparison_mode:
    st.session_state.comparison_mode = comparison_mode
    st.rerun()

# Configurable thresholds
st.sidebar.subheader("Performance Thresholds")
strong_threshold = st.sidebar.number_input("Strong Performance Threshold (%)", 
                                         min_value=0.0, max_value=100.0, 
                                         value=STRONG_TH, step=1.0)
average_threshold = st.sidebar.number_input("Average Performance Threshold (%)", 
                                          min_value=0.0, max_value=100.0, 
                                          value=AVERAGE_LO, step=1.0)
top_count = st.sidebar.number_input("Number of Top Students to Show", 
                                  min_value=1, max_value=20, value=TOP_COUNT, step=1)
bottom_count = st.sidebar.number_input("Number of Bottom Students to Show", 
                                     min_value=1, max_value=20, value=BOTTOM_COUNT, step=1)

# Logo upload
show_logo = st.sidebar.checkbox("Upload logo to include in reports", value=False)
logo_file = None
if show_logo:
    logo_file = st.sidebar.file_uploader("Upload logo image (png/jpg)", type=["png", "jpg", "jpeg"])

report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
st.sidebar.markdown(f"**Report date:** {report_date}")

# ---------- Enhanced Parsing functions ----------
def find_fraction(s: str):
    """Extract numerator and denominator from fraction strings"""
    if not isinstance(s, str):
        return None
    m = re.search(r"(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?)", s)
    if m:
        try:
            numerator = float(m.group(1))
            denominator = float(m.group(2))
            if denominator == 0:
                return None
            return numerator, denominator
        except (ValueError, TypeError):
            return None
    return None

def find_number(s: str):
    """Extract numeric values from strings"""
    if not isinstance(s, str):
        return None
    m = re.search(r"(\d+(?:\.\d+)?)", s)
    if m:
        try:
            return float(m.group(1))
        except (ValueError, TypeError):
            return None
    return None

def parse_series_to_pct(series: pd.Series):
    """Convert various score formats to percentages with enhanced error handling"""
    s = series.copy()
    parsed = []
    orig = s.fillna("").astype(str)
    has_slash = orig.str.contains("/").any()

    if not has_slash:
        numeric = pd.to_numeric(s, errors="coerce")
        nonnull = numeric.dropna()
        
        if len(nonnull) == 0:
            for val in s:
                if pd.isna(val):
                    parsed.append(np.nan)
                    continue
                num = find_number(str(val))
                parsed.append(num if num is not None else np.nan)
            
            parsed = pd.Series(parsed, index=s.index)
            if parsed.dropna().empty:
                return parsed
            maxv = parsed.max()
            if maxv > 100:
                return (parsed / maxv * 100).clip(0, 100)
            return parsed.clip(0, 100)
        else:
            maxv = nonnull.max()
            if maxv <= 100:
                return numeric.clip(0, 100)
            else:
                return (numeric / maxv * 100).clip(0, 100)

    for val in s:
        if pd.isna(val):
            parsed.append(np.nan)
            continue
            
        if isinstance(val, (int, float, np.integer, np.floating)):
            parsed.append(float(val))
            continue
            
        fs = find_fraction(str(val))
        if fs:
            obt, tot = fs
            parsed.append((obt / tot) * 100)
            continue
            
        num = find_number(str(val))
        parsed.append(num if num is not None else np.nan)

    parsed = pd.Series(parsed, index=s.index).astype(float)
    nonnull = parsed.dropna()
    
    if nonnull.empty:
        return parsed
    
    maxv = nonnull.max()
    if maxv > 100:
        return (parsed / maxv * 100).clip(0, 100)
    return parsed.clip(0, 100)

def process_dataframe(df, subjects):
    """Process a dataframe and return processed data with statistics"""
    df_processed = df.copy()
    df_processed.columns = df_processed.columns.str.strip()
    
    # Apply parsing to subjects
    pct_cols = []
    for subj in subjects:
        pct_col = subj + "_pct"
        df_processed[pct_col] = parse_series_to_pct(df_processed[subj])
        pct_cols.append(pct_col)

    # Calculate overall percentage
    if "Current Percentage" in df_processed.columns:
        current_pct_numeric = pd.to_numeric(df_processed["Current Percentage"], errors="coerce")
        if current_pct_numeric.notna().sum() > len(df_processed) * 0.5:
            df_processed["Overall_pct"] = current_pct_numeric
        else:
            df_processed["Overall_pct"] = df_processed[pct_cols].mean(axis=1)
    else:
        df_processed["Overall_pct"] = df_processed[pct_cols].mean(axis=1)

    # Remove students with no valid percentage data
    initial_count = len(df_processed)
    df_processed = df_processed[df_processed["Overall_pct"].notna()]
    
    # Calculate statistics
    total_students = len(df_processed)
    highest_pct = df_processed["Overall_pct"].max(skipna=True)
    lowest_pct = df_processed["Overall_pct"].min(skipna=True)
    class_avg = df_processed["Overall_pct"].mean(skipna=True)
    class_median = df_processed["Overall_pct"].median(skipna=True)

    df_sorted = df_processed.sort_values("Overall_pct", ascending=False).reset_index(drop=True)
    top_students = df_sorted.head(top_count)
    
    # Enhanced lowest students logic
    nonzero_df = df_sorted[df_sorted["Overall_pct"] > 0].copy()
    if len(nonzero_df) >= bottom_count:
        bottom_students = nonzero_df.tail(bottom_count).sort_values("Overall_pct", ascending=True)
    else:
        bottom_students = nonzero_df.sort_values("Overall_pct", ascending=True)

    # Performance categories
    strong_students = df_processed[df_processed["Overall_pct"] >= strong_threshold]
    avg_students = df_processed[(df_processed["Overall_pct"] >= average_threshold) & 
                               (df_processed["Overall_pct"] < strong_threshold)]
    weak_students = df_processed[df_processed["Overall_pct"] < average_threshold]

    # Subject averages
    subject_stats = {}
    for subject in subjects:
        pct_col = subject + "_pct"
        if pct_col in df_processed.columns:
            subject_data = df_processed[pct_col].dropna()
            if len(subject_data) > 0:
                subject_stats[subject] = {
                    'average': subject_data.mean(),
                    'median': subject_data.median(),
                    'highest': subject_data.max(),
                    'lowest': subject_data.min(),
                    'std_dev': subject_data.std()
                }

    return {
        'df': df_processed,
        'total_students': total_students,
        'highest_pct': highest_pct,
        'lowest_pct': lowest_pct,
        'class_avg': class_avg,
        'class_median': class_median,
        'top_students': top_students,
        'bottom_students': bottom_students,
        'strong_students': strong_students,
        'avg_students': avg_students,
        'weak_students': weak_students,
        'subject_stats': subject_stats,
        'pct_cols': pct_cols,
        'subjects': subjects
    }

# ---------- Individual Student Analysis Functions ----------
def create_student_performance_chart(student_data, subjects):
    """Create radar chart for individual student performance"""
    # Get student scores for each subject
    scores = []
    for subject in subjects:
        pct_col = subject + "_pct"
        score = student_data.get(pct_col, 0)
        scores.append(score if not pd.isna(score) else 0)
    
    # Number of variables
    num_vars = len(subjects)
    
    # Compute angle for each subject
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    angles += angles[:1]  # Complete the circle
    scores += scores[:1]  # Complete the circle
    
    # Create the plot
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
    
    # Plot the student's scores
    ax.plot(angles, scores, 'o-', linewidth=2, label='Student Score', color='blue')
    ax.fill(angles, scores, alpha=0.25, color='blue')
    
    # Plot class average (as a reference circle at 50%)
    class_avg_line = [50] * (num_vars + 1)
    ax.plot(angles, class_avg_line, '--', linewidth=1, label='Class Average (50%)', color='red', alpha=0.7)
    
    # Add subject labels
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(subjects)
    
    # Set y-axis limits and labels
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'])
    
    # Add title and legend
    student_name = student_data.get('Student Name', 'Unknown Student')
    ax.set_title(f'Subject Performance Radar - {student_name}', size=14, fontweight='bold', pad=20)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.0))
    
    plt.tight_layout()
    return fig_to_bytes(fig)

def create_student_subject_bar_chart(student_data, subjects, class_avg_data):
    """Create bar chart comparing student performance to class average"""
    student_scores = []
    class_avgs = []
    
    for subject in subjects:
        pct_col = subject + "_pct"
        student_score = student_data.get(pct_col, 0)
        student_scores.append(student_score if not pd.isna(student_score) else 0)
        class_avgs.append(class_avg_data.get(subject, {}).get('average', 0))
    
    x = np.arange(len(subjects))
    width = 0.35
    
    fig, ax = plt.subplots(figsize=(12, 6))
    bars1 = ax.bar(x - width/2, student_scores, width, label='Student Score', color='#3498db', alpha=0.8)
    bars2 = ax.bar(x + width/2, class_avgs, width, label='Class Average', color='#e74c3c', alpha=0.6)
    
    ax.set_xlabel('Subjects', fontweight='bold')
    ax.set_ylabel('Percentage (%)', fontweight='bold')
    student_name = student_data.get('Student Name', 'Unknown Student')
    ax.set_title(f'{student_name} - Subject Performance vs Class Average', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(subjects, rotation=45, ha='right')
    ax.legend()
    ax.grid(axis='y', alpha=0.3)
    ax.set_ylim(0, 100)
    
    # Add value labels on bars
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                   f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
    
    fig.tight_layout()
    return fig_to_bytes(fig)

def get_student_performance_insights(student_data, subjects, class_avg_data):
    """Generate personalized insights for a student"""
    insights = []
    student_name = student_data.get('Student Name', 'Unknown Student')
    overall_pct = student_data.get('Overall_pct', 0)
    
    # Overall performance insight
    if overall_pct >= strong_threshold:
        insights.append(f"🎉 **Excellent Performance**: {student_name} is performing strongly with an overall score of {overall_pct:.1f}%")
    elif overall_pct >= average_threshold:
        insights.append(f"📊 **Solid Performance**: {student_name} is performing at an average level with {overall_pct:.1f}%")
    else:
        insights.append(f"📚 **Needs Improvement**: {student_name} needs additional support with an overall score of {overall_pct:.1f}%")
    
    # Subject-specific insights
    strong_subjects = []
    weak_subjects = []
    
    for subject in subjects:
        pct_col = subject + "_pct"
        student_score = student_data.get(pct_col, 0)
        class_avg = class_avg_data.get(subject, {}).get('average', 0)
        
        if not pd.isna(student_score):
            if student_score >= 80:
                strong_subjects.append(subject)
            elif student_score <= 40:
                weak_subjects.append(subject)
            elif student_score < class_avg - 10:
                weak_subjects.append(f"{subject} ({student_score:.1f}% vs class avg {class_avg:.1f}%)")
    
    if strong_subjects:
        insights.append(f"🌟 **Strengths**: Excels in {', '.join(strong_subjects)}")
    if weak_subjects:
        insights.append(f"📉 **Areas for Improvement**: Needs support in {', '.join(weak_subjects)}")
    
    # Rank insight if available
    if 'Current Position' in student_data and not pd.isna(student_data['Current Position']):
        total_students = len(class_avg_data.get('all_students', []))
        if total_students > 0:
            rank = student_data['Current Position']
            percentile = (rank / total_students) * 100
            insights.append(f"🏆 **Class Rank**: Position {rank} out of {total_students} students ({percentile:.1f} percentile)")
    
    return insights

# ---------- Comparison Functions ----------
def compare_weeks(week1_data, week2_data, week1_name="Week 1", week2_name="Week 2"):
    """Compare two weeks of data and return comparison metrics"""
    
    comparison = {
        'class_avg_change': week2_data['class_avg'] - week1_data['class_avg'],
        'class_avg_change_pct': ((week2_data['class_avg'] - week1_data['class_avg']) / week1_data['class_avg'] * 100) if week1_data['class_avg'] > 0 else 0,
        'highest_change': week2_data['highest_pct'] - week1_data['highest_pct'],
        'lowest_change': week2_data['lowest_pct'] - week1_data['lowest_pct'],
        'strong_students_change': len(week2_data['strong_students']) - len(week1_data['strong_students']),
        'weak_students_change': len(week2_data['weak_students']) - len(week1_data['weak_students']),
        'subject_comparison': {}
    }
    
    # Compare subject averages
    for subject in week1_data['subjects']:
        if subject in week1_data['subject_stats'] and subject in week2_data['subject_stats']:
            week1_avg = week1_data['subject_stats'][subject]['average']
            week2_avg = week2_data['subject_stats'][subject]['average']
            change = week2_avg - week1_avg
            change_pct = (change / week1_avg * 100) if week1_avg > 0 else 0
            
            comparison['subject_comparison'][subject] = {
                'week1_avg': week1_avg,
                'week2_avg': week2_avg,
                'change': change,
                'change_pct': change_pct
            }
    
    # Student-level comparison (for students present in both weeks)
    common_students = []
    student_subject_comparison = {}
    
    if 'Student Name' in week1_data['df'].columns and 'Student Name' in week2_data['df'].columns:
        week1_students = set(week1_data['df']['Student Name'].dropna())
        week2_students = set(week2_data['df']['Student Name'].dropna())
        common_names = week1_students.intersection(week2_students)
        
        for name in common_names:
            week1_row = week1_data['df'][week1_data['df']['Student Name'] == name].iloc[0]
            week2_row = week2_data['df'][week2_data['df']['Student Name'] == name].iloc[0]
            
            week1_score = week1_row['Overall_pct']
            week2_score = week2_row['Overall_pct']
            change = week2_score - week1_score
            
            # Subject-wise comparison for this student
            subject_changes = {}
            for subject in week1_data['subjects']:
                week1_subj = week1_row.get(subject + '_pct', np.nan)
                week2_subj = week2_row.get(subject + '_pct', np.nan)
                if not pd.isna(week1_subj) and not pd.isna(week2_subj):
                    subject_changes[subject] = {
                        'week1_score': week1_subj,
                        'week2_score': week2_subj,
                        'change': week2_subj - week1_subj
                    }
            
            common_students.append({
                'name': name,
                'week1_score': week1_score,
                'week2_score': week2_score,
                'change': change,
                'change_pct': (change / week1_score * 100) if week1_score > 0 else 0,
                'subject_changes': subject_changes
            })
            
            student_subject_comparison[name] = subject_changes
    
    comparison['common_students'] = pd.DataFrame(common_students)
    comparison['common_students_count'] = len(common_students)
    comparison['student_subject_comparison'] = student_subject_comparison
    
    return comparison

# ---------- Visualization Functions ----------
def fig_to_bytes(fig, dpi=150):
    """Convert matplotlib figure to bytes"""
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=dpi)
    plt.close(fig)
    buf.seek(0)
    return buf

def create_subject_avg_bar_chart(subjects_list, df_local):
    """Create bar chart of subject averages"""
    avgs = [df_local[s + "_pct"].mean(skipna=True) for s in subjects_list]
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(subjects_list, avgs, color='skyblue', alpha=0.7)
    ax.set_title("Subject-wise Average Performance (%)", fontsize=14, fontweight='bold')
    ax.set_ylabel("Average %", fontweight='bold')
    ax.set_ylim(0, 100)
    ax.grid(axis='y', alpha=0.3)
    
    for bar, avg in zip(bars, avgs):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                f'{avg:.1f}%', ha='center', va='bottom', fontweight='bold')
    
    ax.set_xticklabels(subjects_list, rotation=45, ha='right')
    fig.tight_layout()
    return fig_to_bytes(fig)

def create_overall_distribution(hist_series):
    """Create histogram of overall percentage distribution"""
    fig, ax = plt.subplots(figsize=(8, 4))
    n, bins, patches = ax.hist(hist_series.dropna(), bins=12, color='lightgreen', 
                              alpha=0.7, edgecolor='black')
    ax.set_title("Overall Percentage Distribution", fontsize=14, fontweight='bold')
    ax.set_xlabel("Percentage", fontweight='bold')
    ax.set_ylabel("Number of Students", fontweight='bold')
    ax.grid(axis='y', alpha=0.3)
    
    for i, (count, patch) in enumerate(zip(n, patches)):
        if count > 0:
            ax.text(patch.get_x() + patch.get_width()/2, count + 0.1,
                   f'{int(count)}', ha='center', va='bottom', fontweight='bold')
    
    fig.tight_layout()
    return fig_to_bytes(fig)

def create_performance_pie_chart(strong_count, avg_count, weak_count):
    """Create pie chart of performance categories"""
    fig, ax = plt.subplots(figsize=(6, 6))
    sizes = [strong_count, avg_count, weak_count]
    labels = [f'Strong\n(≥{strong_threshold}%)', 
             f'Average\n({average_threshold}-{strong_threshold}%)', 
             f'Weak\n(<{average_threshold}%)']
    colors = ['#2ecc71', '#f39c12', '#e74c3c']
    
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%',
                                     startangle=90)
    
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
    
    ax.set_title('Performance Category Distribution', fontsize=14, fontweight='bold')
    return fig_to_bytes(fig)

def create_comparison_chart(week1_data, week2_data, week1_name="Week 1", week2_name="Week 2"):
    """Create comparison chart between two weeks"""
    metrics = ['Class Average', 'Highest Score', 'Lowest Score']
    week1_values = [week1_data['class_avg'], week1_data['highest_pct'], week1_data['lowest_pct']]
    week2_values = [week2_data['class_avg'], week2_data['highest_pct'], week2_data['lowest_pct']]
    
    x = np.arange(len(metrics))
    width = 0.35
    
    fig, ax = plt.subplots(figsize=(10, 6))
    bars1 = ax.bar(x - width/2, week1_values, width, label=week1_name, color='#3498db', alpha=0.7)
    bars2 = ax.bar(x + width/2, week2_values, width, label=week2_name, color='#2ecc71', alpha=0.7)
    
    ax.set_ylabel('Percentage', fontweight='bold')
    ax.set_title('Week-over-Week Comparison', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(metrics)
    ax.legend()
    ax.grid(axis='y', alpha=0.3)
    
    # Add value labels on bars
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                   f'{height:.1f}%', ha='center', va='bottom', fontweight='bold')
    
    fig.tight_layout()
    return fig_to_bytes(fig)

def create_subject_comparison_chart(comparison_data, week1_name="Week 1", week2_name="Week 2"):
    """Create subject-wise comparison chart"""
    subjects = list(comparison_data.keys())
    week1_avgs = [comparison_data[subj]['week1_avg'] for subj in subjects]
    week2_avgs = [comparison_data[subj]['week2_avg'] for subj in subjects]
    changes = [comparison_data[subj]['change'] for subj in subjects]
    
    x = np.arange(len(subjects))
    width = 0.35
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 10))
    
    # Bar chart for averages
    bars1 = ax1.bar(x - width/2, week1_avgs, width, label=week1_name, color='#3498db', alpha=0.7)
    bars2 = ax1.bar(x + width/2, week2_avgs, width, label=week2_name, color='#2ecc71', alpha=0.7)
    
    ax1.set_ylabel('Average %', fontweight='bold')
    ax1.set_title('Subject-wise Average Comparison', fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(subjects, rotation=45, ha='right')
    ax1.legend()
    ax1.grid(axis='y', alpha=0.3)
    
    # Add value labels
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + 1,
                    f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
    
    # Bar chart for changes
    colors = ['#2ecc71' if x >= 0 else '#e74c3c' for x in changes]
    bars3 = ax2.bar(x, changes, color=colors, alpha=0.7)
    ax2.set_ylabel('Change (%)', fontweight='bold')
    ax2.set_title('Subject-wise Performance Change', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(subjects, rotation=45, ha='right')
    ax2.grid(axis='y', alpha=0.3)
    
    # Add change labels
    for bar, change in zip(bars3, changes):
        height = bar.get_height()
        va = 'bottom' if height >= 0 else 'top'
        color = 'green' if height >= 0 else 'red'
        ax2.text(bar.get_x() + bar.get_width()/2., height + (0.5 if height >= 0 else -0.5),
                f'{change:+.1f}%', ha='center', va=va, color=color, fontweight='bold')
    
    ax2.axhline(y=0, color='black', linestyle='-', alpha=0.3)
    
    fig.tight_layout()
    return fig_to_bytes(fig)

def create_student_comparison_chart(student_data, week1_name="Week 1", week2_name="Week 2"):
    """Create individual student subject comparison chart"""
    subjects = list(student_data.keys())
    week1_scores = [student_data[subj]['week1_score'] for subj in subjects]
    week2_scores = [student_data[subj]['week2_score'] for subj in subjects]
    changes = [student_data[subj]['change'] for subj in subjects]
    
    x = np.arange(len(subjects))
    width = 0.35
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 8))
    
    # Bar chart for scores
    bars1 = ax1.bar(x - width/2, week1_scores, width, label=week1_name, color='#3498db', alpha=0.7)
    bars2 = ax1.bar(x + width/2, week2_scores, width, label=week2_name, color='#2ecc71', alpha=0.7)
    
    ax1.set_ylabel('Score (%)', fontweight='bold')
    ax1.set_title('Student Subject Performance Comparison', fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(subjects, rotation=45, ha='right')
    ax1.legend()
    ax1.grid(axis='y', alpha=0.3)
    ax1.set_ylim(0, 100)
    
    # Add value labels
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + 1,
                    f'{height:.1f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
    
    # Bar chart for changes
    colors = ['#2ecc71' if x >= 0 else '#e74c3c' for x in changes]
    bars3 = ax2.bar(x, changes, color=colors, alpha=0.7)
    ax2.set_ylabel('Change (%)', fontweight='bold')
    ax2.set_title('Subject-wise Performance Change', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(subjects, rotation=45, ha='right')
    ax2.grid(axis='y', alpha=0.3)
    
    # Add change labels
    for bar, change in zip(bars3, changes):
        height = bar.get_height()
        va = 'bottom' if height >= 0 else 'top'
        color = 'green' if height >= 0 else 'red'
        ax2.text(bar.get_x() + bar.get_width()/2., height + (0.5 if height >= 0 else -0.5),
                f'{change:+.1f}%', ha='center', va=va, color=color, fontweight='bold')
    
    ax2.axhline(y=0, color='black', linestyle='-', alpha=0.3)
    
    fig.tight_layout()
    return fig_to_bytes(fig)

# ---------- PDF Generation for Single File ----------
def generate_pdf_buffer(df_processed, subjects, logo_bytes=None):
    """Generate PDF report buffer for single file analysis"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []

    # Header with logo
    header_data = []
    if logo_bytes:
        try:
            img = utils.ImageReader(logo_bytes)
            iw, ih = img.getSize()
            aspect = ih / float(iw)
            rl_img = RLImage(logo_bytes, width=80, height=(80 * aspect))
            header_data.append([rl_img, Paragraph(f"<b>{academy_name}</b><br/><i>Student Performance Report</i><br/>{report_date}", styles["Normal"])])
            t = Table(header_data, colWidths=[90, 420])
            t.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
            story.append(t)
        except Exception as e:
            story.append(Paragraph(f"<b>{academy_name}</b>", styles["Title"]))
            story.append(Paragraph("Student Performance Report", styles["Heading2"]))
    else:
        story.append(Paragraph(f"<b>{academy_name}</b>", styles["Title"]))
        story.append(Paragraph("Student Performance Report", styles["Heading2"]))
    
    story.append(Paragraph(f"Report generated: {report_date}", styles["Normal"]))
    story.append(Spacer(1, 12))

    # Class Information
    if "Class Name" in df_processed.columns:
        class_name = df_processed["Class Name"].iloc[0] if len(df_processed) > 0 else "Unknown"
        story.append(Paragraph(f"<b>Class:</b> {class_name}", styles["Normal"]))
    
    if "Section Name" in df_processed.columns:
        section_name = df_processed["Section Name"].iloc[0] if len(df_processed) > 0 else "Unknown"
        story.append(Paragraph(f"<b>Section:</b> {section_name}", styles["Normal"]))
    
    story.append(Spacer(1, 12))

    # Key Statistics
    story.append(Paragraph("<b>Key Statistics</b>", styles["Heading2"]))
    
    total_students = len(df_processed)
    class_avg = df_processed["Overall_pct"].mean()
    highest_pct = df_processed["Overall_pct"].max()
    lowest_pct = df_processed["Overall_pct"].min()
    class_median = df_processed["Overall_pct"].median()
    
    # Performance categories
    strong_students = df_processed[df_processed["Overall_pct"] >= strong_threshold]
    avg_students = df_processed[(df_processed["Overall_pct"] >= average_threshold) & (df_processed["Overall_pct"] < strong_threshold)]
    weak_students = df_processed[df_processed["Overall_pct"] < average_threshold]
    
    strong_count = len(strong_students)
    avg_count = len(avg_students)
    weak_count = len(weak_students)
    
    stats_data = [
        ["Metric", "Value"],
        ["Total Students", str(total_students)],
        ["Class Average", f"{class_avg:.2f}%"],
        ["Class Median", f"{class_median:.2f}%"],
        ["Highest Score", f"{highest_pct:.2f}%"],
        ["Lowest Score", f"{lowest_pct:.2f}%"],
        [f"Strong Students (≥{strong_threshold}%)", f"{strong_count} ({(strong_count/total_students)*100:.1f}%)"],
        [f"Average Students ({average_threshold}-{strong_threshold}%)", f"{avg_count} ({(avg_count/total_students)*100:.1f}%)"],
        [f"Weak Students (<{average_threshold}%)", f"{weak_count} ({(weak_count/total_students)*100:.1f}%)"]
    ]
    
    stats_table = Table(stats_data, colWidths=[200, 100])
    stats_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold")
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 12))

    # Subject-wise Performance
    story.append(Paragraph("<b>Subject-wise Performance</b>", styles["Heading2"]))
    
    subject_data = [["Subject", "Average", "Median", "Highest", "Lowest", "Std Dev"]]
    for subject in subjects:
        pct_col = subject + "_pct"
        if pct_col in df_processed.columns:
            subject_pct = df_processed[pct_col].dropna()
            if len(subject_pct) > 0:
                subject_data.append([
                    subject,
                    f"{subject_pct.mean():.2f}%",
                    f"{subject_pct.median():.2f}%",
                    f"{subject_pct.max():.2f}%",
                    f"{subject_pct.min():.2f}%",
                    f"{subject_pct.std():.2f}%"
                ])
    
    subject_table = Table(subject_data, colWidths=[120, 70, 70, 70, 70, 70])
    subject_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
    ]))
    story.append(subject_table)
    story.append(PageBreak())

    # Top Students
    story.append(Paragraph("<b>Top Performing Students</b>", styles["Heading2"]))
    
    top_students = df_processed.nlargest(top_count, "Overall_pct")
    top_data = [["Rank", "Student Name", "Overall %"] + [f"{s} %" for s in subjects]]
    
    for i, (idx, student) in enumerate(top_students.iterrows(), 1):
        row = [str(i), student.get("Student Name", "Unknown")]
        row.append(f"{student['Overall_pct']:.2f}%")
        
        for subject in subjects:
            pct_col = subject + "_pct"
            score = student.get(pct_col, np.nan)
            row.append(f"{score:.2f}%" if not pd.isna(score) else "N/A")
        
        top_data.append(row)
    
    # Adjust column widths based on number of subjects
    col_widths = [40, 120, 60] + [50] * len(subjects)
    top_table = Table(top_data, colWidths=col_widths)
    top_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 7),
    ]))
    story.append(top_table)
    story.append(Spacer(1, 12))

    # Average Performing Students
    story.append(Paragraph("<b>Average Performing Students</b>", styles["Heading2"]))
    story.append(Paragraph(f"Students scoring between {average_threshold}% and {strong_threshold}%", styles["Normal"]))
    
    if len(avg_students) > 0:
        # Sort average students by overall percentage (descending)
        avg_students_sorted = avg_students.sort_values("Overall_pct", ascending=False)
        
        avg_data = [["Rank", "Student Name", "Overall %"] + [f"{s} %" for s in subjects]]
        
        for i, (idx, student) in enumerate(avg_students_sorted.iterrows(), 1):
            row = [str(i), student.get("Student Name", "Unknown")]
            row.append(f"{student['Overall_pct']:.2f}%")
            
            for subject in subjects:
                pct_col = subject + "_pct"
                score = student.get(pct_col, np.nan)
                row.append(f"{score:.2f}%" if not pd.isna(score) else "N/A")
            
            avg_data.append(row)
        
        avg_table = Table(avg_data, colWidths=col_widths)
        avg_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 7),
            ("BACKGROUND", (0,1), (-1,-1), colors.lightyellow),  # Light yellow for average students
        ]))
        story.append(avg_table)
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"Total Average Students: {len(avg_students_sorted)} ({(len(avg_students_sorted)/total_students)*100:.1f}% of class)", styles["Normal"]))
    else:
        story.append(Paragraph("No students found in the average performance category.", styles["Normal"]))
    
    story.append(Spacer(1, 12))

    # Bottom Students (excluding zeros)
    story.append(Paragraph("<b>Students Needing Attention</b>", styles["Heading2"]))
    
    non_zero = df_processed[df_processed["Overall_pct"] > 0]
    if len(non_zero) >= bottom_count:
        bottom_students = non_zero.nsmallest(bottom_count, "Overall_pct")
    else:
        bottom_students = non_zero.nsmallest(len(non_zero), "Overall_pct")
    
    bottom_data = [["Rank", "Student Name", "Overall %"] + [f"{s} %" for s in subjects]]
    
    for i, (idx, student) in enumerate(bottom_students.iterrows(), 1):
        row = [str(i), student.get("Student Name", "Unknown")]
        row.append(f"{student['Overall_pct']:.2f}%")
        
        for subject in subjects:
            pct_col = subject + "_pct"
            score = student.get(pct_col, np.nan)
            row.append(f"{score:.2f}%" if not pd.isna(score) else "N/A")
        
        bottom_data.append(row)
    
    bottom_table = Table(bottom_data, colWidths=col_widths)
    bottom_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 7),
        ("TEXTCOLOR", (0,1), (-1,-1), colors.red),  # Red color for bottom students
    ]))
    story.append(bottom_table)
    story.append(PageBreak())

    # NEW: Individual Student Analysis Section
    story.append(Paragraph("<b>Individual Student Detailed Analysis</b>", styles["Heading2"]))
    story.append(Paragraph("Comprehensive performance breakdown for each student", styles["Normal"]))
    story.append(Spacer(1, 12))

    # Calculate subject statistics for comparison
    subject_stats = {}
    for subject in subjects:
        pct_col = subject + "_pct"
        if pct_col in df_processed.columns:
            subject_data = df_processed[pct_col].dropna()
            if len(subject_data) > 0:
                subject_stats[subject] = {
                    'average': subject_data.mean(),
                    'median': subject_data.median(),
                    'highest': subject_data.max(),
                    'lowest': subject_data.min(),
                    'std_dev': subject_data.std()
                }

    # Sort students by overall percentage for organized presentation
    sorted_students = df_processed.sort_values("Overall_pct", ascending=False)
    
    students_per_page = 2  # Adjust based on content density
    student_count = 0
    
    for idx, student in sorted_students.iterrows():
        if student_count > 0 and student_count % students_per_page == 0:
            story.append(PageBreak())
        
        student_name = student.get("Student Name", "Unknown Student")
        overall_pct = student.get("Overall_pct", 0)
        
        # Student Header
        story.append(Paragraph(f"<b>{student_name}</b>", styles["Heading3"]))
        
        # Student Overview Table
        overview_data = [
            ["Metric", "Value", "Performance Category"],
            ["Overall Percentage", f"{overall_pct:.2f}%", 
             "Strong" if overall_pct >= strong_threshold else 
             "Average" if overall_pct >= average_threshold else "Weak"],
            ["Class Rank", f"{idx + 1} / {len(sorted_students)}", 
             f"Top {(idx + 1)/len(sorted_students)*100:.1f}%"]
        ]
        
        if 'Current Position' in student and not pd.isna(student['Current Position']):
            overview_data[2][1] = f"{student['Current Position']} / {len(sorted_students)}"
        
        overview_table = Table(overview_data, colWidths=[150, 100, 150])
        overview_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (2,1), (2,1), 
             colors.lightgreen if overall_pct >= strong_threshold else 
             colors.lightyellow if overall_pct >= average_threshold else colors.pink),
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 8))

        # Subject Performance Table
        story.append(Paragraph("<b>Subject-wise Performance</b>", styles["Heading4"]))
        
        subject_performance_data = [["Subject", "Student Score", "Class Average", "Difference", "Status"]]
        
        strong_subjects = []
        weak_subjects = []
        
        for subject in subjects:
            pct_col = subject + "_pct"
            student_score = student.get(pct_col, np.nan)
            class_avg = subject_stats.get(subject, {}).get('average', np.nan)
            
            if not pd.isna(student_score) and not pd.isna(class_avg):
                difference = student_score - class_avg
                status = "Above Average" if difference > 5 else "Below Average" if difference < -5 else "At Average"
                
                # Color coding for status
                if difference > 10:
                    strong_subjects.append(subject)
                elif difference < -10:
                    weak_subjects.append(subject)
                
                subject_performance_data.append([
                    subject,
                    f"{student_score:.2f}%",
                    f"{class_avg:.2f}%",
                    f"{difference:+.2f}%",
                    status
                ])
        
        subject_perf_table = Table(subject_performance_data, colWidths=[100, 80, 80, 80, 100])
        subject_perf_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 7),
        ]))
        
        # Apply color to status column
        for i in range(1, len(subject_performance_data)):
            status = subject_performance_data[i][4]
            if status == "Above Average":
                subject_perf_table.setStyle(TableStyle([
                    ("TEXTCOLOR", (4,i), (4,i), colors.darkgreen),
                    ("FONTNAME", (4,i), (4,i), "Helvetica-Bold")
                ]))
            elif status == "Below Average":
                subject_perf_table.setStyle(TableStyle([
                    ("TEXTCOLOR", (4,i), (4,i), colors.darkred),
                    ("FONTNAME", (4,i), (4,i), "Helvetica-Bold")
                ]))
        
        story.append(subject_perf_table)
        story.append(Spacer(1, 8))

        # Performance Insights - FIXED VERSION (no emojis, proper HTML)
        story.append(Paragraph("<b>Performance Insights</b>", styles["Heading4"]))
        
        insights = []
        
        # Overall performance insight
        if overall_pct >= strong_threshold:
            insights.append(f"<b>Excellent Performance</b>: {student_name} is performing strongly with an overall score of {overall_pct:.1f}%")
        elif overall_pct >= average_threshold:
            insights.append(f"<b>Solid Performance</b>: {student_name} is performing at an average level with {overall_pct:.1f}%")
        else:
            insights.append(f"<b>Needs Improvement</b>: {student_name} needs additional support with an overall score of {overall_pct:.1f}%")
        
        # Subject-specific insights
        if strong_subjects:
            insights.append(f"<b>Strengths</b>: Excels in {', '.join(strong_subjects)}")
        
        if weak_subjects:
            insights.append(f"<b>Areas for Improvement</b>: Needs support in {', '.join(weak_subjects)}")
        
        # Add recommendations based on performance
        if overall_pct >= strong_threshold:
            insights.append("<b>Recommendation</b>: Continue current study habits and consider advanced topics")
        elif overall_pct >= average_threshold:
            insights.append("<b>Recommendation</b>: Focus on weak subjects to reach strong performance level")
        else:
            insights.append("<b>Recommendation</b>: Requires intensive support and regular progress monitoring")
        
        # Add insights to story
        for insight in insights:
            story.append(Paragraph(insight, styles["Normal"]))
            story.append(Spacer(1, 4))
        
        story.append(Spacer(1, 12))
        student_count += 1

    doc.build(story)
    buffer.seek(0)
    return buffer
def generate_docx_buffer(df_processed, subjects, logo_bytes=None):
    """Generate Word document report buffer for single file analysis"""
    # This is a placeholder - implement Word document generation here
    document = Document()
    
    # Add basic content
    document.add_heading(f'{academy_name} - Student Performance Report', 0)
    document.add_paragraph(f'Report generated: {report_date}')
    
    # Save to buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- PDF Generation for Comparison ----------
def generate_comparison_pdf_buffer(week1_data, week2_data, comparison_data, week1_name, week2_name, logo_bytes=None):
    """Generate PDF report buffer for comparison"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []

    # Header with logo
    header_data = []
    if logo_bytes:
        try:
            img = utils.ImageReader(logo_bytes)
            iw, ih = img.getSize()
            aspect = ih / float(iw)
            rl_img = RLImage(logo_bytes, width=80, height=(80 * aspect))
            header_data.append([rl_img, Paragraph(f"<b>{academy_name}</b><br/><i>Two-Week Performance Comparison</i><br/>{report_date}", styles["Normal"])])
            t = Table(header_data, colWidths=[90, 420])
            t.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
            story.append(t)
        except Exception as e:
            story.append(Paragraph(f"<b>{academy_name}</b>", styles["Title"]))
            story.append(Paragraph("Two-Week Performance Comparison", styles["Heading2"]))
    else:
        story.append(Paragraph(f"<b>{academy_name}</b>", styles["Title"]))
        story.append(Paragraph("Two-Week Performance Comparison", styles["Heading2"]))
    
    story.append(Paragraph(f"Report generated: {report_date}", styles["Normal"]))
    story.append(Spacer(1, 12))

    # Comparison Periods
    story.append(Paragraph("<b>Comparison Periods</b>", styles["Heading3"]))
    period_data = [
        ["Period", "Name", "Total Students", "Class Average"],
        [week1_name, week1_name, str(week1_data['total_students']), f"{week1_data['class_avg']:.2f}%"],
        [week2_name, week2_name, str(week2_data['total_students']), f"{week2_data['class_avg']:.2f}%"]
    ]
    period_table = Table(period_data, colWidths=[100, 150, 100, 100])
    period_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold")
    ]))
    story.append(period_table)
    story.append(Spacer(1,12))

    # Key Metrics Comparison
    story.append(Paragraph("<b>Key Metrics Comparison</b>", styles["Heading2"]))
    metrics_data = [
        ["Metric", f"{week1_name}", f"{week2_name}", "Change", "Change %"],
        ["Class Average", f"{week1_data['class_avg']:.2f}%", f"{week2_data['class_avg']:.2f}%", 
         f"{comparison_data['class_avg_change']:+.2f}%", f"{comparison_data['class_avg_change_pct']:+.1f}%"],
        ["Highest Score", f"{week1_data['highest_pct']:.2f}%", f"{week2_data['highest_pct']:.2f}%", 
         f"{comparison_data['highest_change']:+.2f}%", "-"],
        ["Lowest Score", f"{week1_data['lowest_pct']:.2f}%", f"{week2_data['lowest_pct']:.2f}%", 
         f"{comparison_data['lowest_change']:+.2f}%", "-"],
        ["Strong Students", str(len(week1_data['strong_students'])), str(len(week2_data['strong_students'])), 
         f"{comparison_data['strong_students_change']:+d}", "-"],
        ["Weak Students", str(len(week1_data['weak_students'])), str(len(week2_data['weak_students'])), 
         f"{comparison_data['weak_students_change']:+d}", "-"]
    ]
    metrics_table = Table(metrics_data, colWidths=[120, 80, 80, 80, 80])
    metrics_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("BACKGROUND", (3,1), (3,-1), colors.lightgreen if comparison_data['class_avg_change'] >= 0 else colors.pink),
        ("BACKGROUND", (4,1), (4,1), colors.lightgreen if comparison_data['class_avg_change'] >= 0 else colors.pink),
    ]))
    story.append(metrics_table)
    story.append(Spacer(1,12))

    # Subject-wise Comparison
    story.append(Paragraph("<b>Subject-wise Performance Comparison</b>", styles["Heading2"]))
    subject_data = [["Subject", f"{week1_name} Avg", f"{week2_name} Avg", "Change", "Change %"]]
    
    for subject, stats in comparison_data['subject_comparison'].items():
        change_color = colors.lightgreen if stats['change'] >= 0 else colors.pink
        subject_data.append([
            subject,
            f"{stats['week1_avg']:.2f}%",
            f"{stats['week2_avg']:.2f}%",
            f"{stats['change']:+.2f}%",
            f"{stats['change_pct']:+.1f}%"
        ])
    
    subject_table = Table(subject_data, colWidths=[120, 80, 80, 80, 80])
    subject_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.3, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("BACKGROUND", (3,1), (4,-1), colors.lightgreen),
    ]))
    
    # Apply color to negative changes
    for i in range(1, len(subject_data)):
        if comparison_data['subject_comparison'][subject_data[i][0]]['change'] < 0:
            subject_table.setStyle(TableStyle([
                ("BACKGROUND", (3,i), (4,i), colors.pink)
            ]))
    
    story.append(subject_table)
    story.append(PageBreak())

    # Individual Student Comparison
    if comparison_data['common_students_count'] > 0:
        story.append(Paragraph("<b>Individual Student Performance Comparison</b>", styles["Heading2"]))
        story.append(Paragraph(f"Tracking {comparison_data['common_students_count']} common students", styles["Normal"]))
        story.append(Spacer(1,12))

        # Top Improvers
        story.append(Paragraph("<b>Top 5 Improvers</b>", styles["Heading3"]))
        top_improvers = comparison_data['common_students'].nlargest(5, 'change')
        improver_data = [["Student", f"{week1_name}", f"{week2_name}", "Change", "Change %"]]
        
        for _, student in top_improvers.iterrows():
            improver_data.append([
                student['name'],
                f"{student['week1_score']:.2f}%",
                f"{student['week2_score']:.2f}%",
                f"{student['change']:+.2f}%",
                f"{student['change_pct']:+.1f}%"
            ])
        
        improver_table = Table(improver_data, colWidths=[200, 80, 80, 80, 80])
        improver_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (3,1), (4,-1), colors.lightgreen),
        ]))
        story.append(improver_table)
        story.append(Spacer(1,12))

        # Top Decliners
        story.append(Paragraph("<b>Top 5 Decliners</b>", styles["Heading3"]))
        top_decliners = comparison_data['common_students'].nsmallest(5, 'change')
        decliner_data = [["Student", f"{week1_name}", f"{week2_name}", "Change", "Change %"]]
        
        for _, student in top_decliners.iterrows():
            decliner_data.append([
                student['name'],
                f"{student['week1_score']:.2f}%",
                f"{student['week2_score']:.2f}%",
                f"{student['change']:+.2f}%",
                f"{student['change_pct']:+.1f}%"
            ])
        
        decliner_table = Table(decliner_data, colWidths=[200, 80, 80, 80, 80])
        decliner_table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.3, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (3,1), (4,-1), colors.pink),
        ]))
        story.append(decliner_table)
        story.append(PageBreak())

        # Detailed Student Subject Analysis
        story.append(Paragraph("<b>Detailed Student Subject Analysis</b>", styles["Heading2"]))
        
        for idx, (student_name, subject_data) in enumerate(comparison_data['student_subject_comparison'].items()):
            if idx > 0 and idx % 2 == 0:
                story.append(PageBreak())
                
            story.append(Paragraph(f"<b>{student_name}</b>", styles["Heading3"]))
            
            student_overall = comparison_data['common_students'][comparison_data['common_students']['name'] == student_name].iloc[0]
            story.append(Paragraph(f"Overall: {student_overall['week1_score']:.2f}% → {student_overall['week2_score']:.2f}% "
                                 f"(Change: {student_overall['change']:+.2f}%)", styles["Normal"]))
            
            # Subject performance table
            subject_rows = [["Subject", f"{week1_name}", f"{week2_name}", "Change"]]
            for subject, scores in subject_data.items():
                subject_rows.append([
                    subject,
                    f"{scores['week1_score']:.2f}%",
                    f"{scores['week2_score']:.2f}%",
                    f"{scores['change']:+.2f}%"
                ])
            
            student_table = Table(subject_rows, colWidths=[150, 80, 80, 80])
            student_table.setStyle(TableStyle([
                ("GRID", (0,0), (-1,-1), 0.3, colors.black),
                ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ]))
            
            # Color code changes
            for i in range(1, len(subject_rows)):
                change = subject_data[subject_rows[i][0]]['change']
                if change >= 0:
                    student_table.setStyle(TableStyle([
                        ("BACKGROUND", (3,i), (3,i), colors.lightgreen)
                    ]))
                else:
                    student_table.setStyle(TableStyle([
                        ("BACKGROUND", (3,i), (3,i), colors.pink)
                    ]))
            
            story.append(student_table)
            story.append(Spacer(1,12))

    # Summary Insights
    story.append(PageBreak())
    story.append(Paragraph("<b>Summary Insights</b>", styles["Heading2"]))
    
    insights = []
    if comparison_data['class_avg_change'] > 0:
        insights.append(f"✅ <b>Overall Improvement</b>: Class average increased by {comparison_data['class_avg_change']:.2f}% ({comparison_data['class_avg_change_pct']:.1f}%)")
    else:
        insights.append(f"⚠️ <b>Overall Decline</b>: Class average decreased by {abs(comparison_data['class_avg_change']):.2f}% ({abs(comparison_data['class_avg_change_pct']):.1f}%)")
    
    if comparison_data['strong_students_change'] > 0:
        insights.append(f"✅ <b>More Strong Performers</b>: Number of strong students increased by {comparison_data['strong_students_change']}")
    else:
        insights.append(f"⚠️ <b>Fewer Strong Performers</b>: Number of strong students decreased by {abs(comparison_data['strong_students_change'])}")
    
    # Find best and worst performing subjects
    if comparison_data['subject_comparison']:
        best_subject = max(comparison_data['subject_comparison'].items(), key=lambda x: x[1]['change'])
        worst_subject = min(comparison_data['subject_comparison'].items(), key=lambda x: x[1]['change'])
        
        insights.append(f"🎯 <b>Best Improvement</b>: {best_subject[0]} improved by {best_subject[1]['change']:+.2f}%")
        insights.append(f"📉 <b>Largest Decline</b>: {worst_subject[0]} declined by {abs(worst_subject[1]['change']):.2f}%")
    
    for insight in insights:
        story.append(Paragraph(insight, styles["Normal"]))
        story.append(Spacer(1,6))

    doc.build(story)
    buffer.seek(0)
    return buffer

# ---------- Main UI ----------
if not st.session_state.comparison_mode:
    # Single file analysis mode
    st.header("📊 Single File Analysis")
    
    uploaded = st.file_uploader("Upload your Excel file (.xlsx)", type=["xlsx"], key="single_file")
    if not uploaded:
        st.info("Upload an Excel file with student records. Example columns: Student Name, Class Name, Chemistry, Biology, Math, English, Urdu, etc.")
        st.stop()

    # Load dataframe
    try:
        with st.spinner("Loading Excel file..."):
            df = pd.read_excel(uploaded)
        st.success("✅ File uploaded and read successfully!")
    except Exception as e:
        st.error(f"❌ Could not read Excel file: {e}")
        st.stop()

    if df.empty:
        st.error("❌ The uploaded Excel file is empty.")
        st.stop()

    df.columns = df.columns.str.strip()
    st.write("📋 Detected columns:", df.columns.tolist())

    # Detect subjects
    ignore_cols = IGNORE_COLS_DEFAULT.copy()
    subjects = [c for c in df.columns if c not in ignore_cols and c.lower() not in [ic.lower() for ic in ignore_cols]]

    if len(subjects) == 0:
        st.error("❌ No subject columns detected. Please check your Excel file column names.")
        st.write("Currently ignored columns:", ignore_cols)
        st.stop()

    st.write("🎯 Detected subjects:", subjects)

    # Student filtering
    st.sidebar.subheader("Student Filtering")
    if "Class Name" in df.columns:
        available_classes = df["Class Name"].unique().tolist()
        selected_class = st.sidebar.selectbox("Filter by Class", ["All Classes"] + available_classes)
        if selected_class != "All Classes":
            df = df[df["Class Name"] == selected_class]

    if "Section Name" in df.columns:
        available_sections = df["Section Name"].unique().tolist()
        selected_section = st.sidebar.selectbox("Filter by Section", ["All Sections"] + available_sections)
        if selected_section != "All Sections":
            df = df[df["Section Name"] == selected_section]

    # Process data
    st.info("🔄 Processing subject scores...")
    progress_bar = st.progress(0)
    
    try:
        data = process_dataframe(df, subjects)
        processed_df = data['df']
        progress_bar.progress(100)
        st.session_state.df_week1 = data  # Store for potential comparison
    except Exception as e:
        st.error(f"❌ Error processing data: {e}")
        st.stop()

    # Display results
    st.markdown("---")
    st.subheader("📈 Performance Overview")
    
    # Key metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Students", data['total_students'])
    with col2:
        st.metric("Class Average", f"{data['class_avg']:.2f}%")
    with col3:
        st.metric("Highest Score", f"{data['highest_pct']:.2f}%")
    with col4:
        st.metric("Lowest Score", f"{data['lowest_pct']:.2f}%")

    # Performance distribution
    st.subheader("🎯 Performance Distribution")
    dist_col1, dist_col2, dist_col3 = st.columns(3)
    with dist_col1:
        st.metric("Strong Students", len(data['strong_students']), 
                 f"{(len(data['strong_students'])/data['total_students'])*100:.1f}%")
    with dist_col2:
        st.metric("Average Students", len(data['avg_students']),
                 f"{(len(data['avg_students'])/data['total_students'])*100:.1f}%")
    with dist_col3:
        st.metric("Weak Students", len(data['weak_students']),
                 f"{(len(data['weak_students'])/data['total_students'])*100:.1f}%")

    # Charts
    st.subheader("📊 Analytics Charts")
    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        st.write("**Subject-wise Averages**")
        subj_img = create_subject_avg_bar_chart(subjects, processed_df)
        st.image(subj_img)
    with chart_col2:
        st.write("**Overall Distribution**")
        dist_img = create_overall_distribution(processed_df["Overall_pct"])
        st.image(dist_img)

    chart_col3, chart_col4 = st.columns(2)
    with chart_col3:
        st.write("**Performance Categories**")
        pie_img = create_performance_pie_chart(len(data['strong_students']), len(data['avg_students']), len(data['weak_students']))
        st.image(pie_img)

    # Top and Bottom Students
    col_left, col_right = st.columns(2)
    with col_left:
        st.write(f"🏅 **Top {top_count} Students**")
        top_display = data['top_students'][['Student Name', 'Overall_pct']].copy()
        top_display["Overall_pct"] = top_display["Overall_pct"].round(2)
        top_display = top_display.rename(columns={"Overall_pct": "Overall %"})
        st.dataframe(top_display, use_container_width=True)

    with col_right:
        st.write(f"📉 **Bottom {bottom_count} Students**")
        bottom_display = data['bottom_students'][['Student Name', 'Overall_pct']].copy()
        bottom_display["Overall_pct"] = bottom_display["Overall_pct"].round(2)
        bottom_display = bottom_display.rename(columns={"Overall_pct": "Overall %"})
        st.dataframe(bottom_display, use_container_width=True)

    # NEW: Average Students Table
    st.subheader("📊 Average Performing Students")
    if len(data['avg_students']) > 0:
        avg_display = data['avg_students'][['Student Name', 'Overall_pct']].copy()
        avg_display["Overall_pct"] = avg_display["Overall_pct"].round(2)
        avg_display = avg_display.rename(columns={"Overall_pct": "Overall %"})
        avg_display = avg_display.sort_values("Overall %", ascending=False)
        st.dataframe(avg_display, use_container_width=True)
        st.info(f"Showing {len(avg_display)} students performing at average level ({average_threshold}-{strong_threshold}%)")
    else:
        st.info("No students found in the average performance category.")

    # NEW: Individual Student Analysis
    st.markdown("---")
    st.subheader("👤 Individual Student Analysis")
    
    if 'Student Name' in processed_df.columns:
        student_names = processed_df['Student Name'].dropna().unique().tolist()
        selected_student = st.selectbox("Select Student for Detailed Analysis", student_names)
        
        if selected_student:
            student_data = processed_df[processed_df['Student Name'] == selected_student].iloc[0]
            
            # Student overview
            col1, col2, col3 = st.columns(3)
            with col1:
                overall_pct = student_data.get('Overall_pct', 0)
                st.metric("Overall Percentage", f"{overall_pct:.2f}%")
            
            with col2:
                # Determine performance category
                if overall_pct >= strong_threshold:
                    category = "Strong"
                    color = "green"
                elif overall_pct >= average_threshold:
                    category = "Average"
                    color = "orange"
                else:
                    category = "Weak"
                    color = "red"
                st.metric("Performance Category", category)
            
            with col3:
                if 'Current Position' in student_data and not pd.isna(student_data['Current Position']):
                    rank = student_data['Current Position']
                    total = len(processed_df)
                    st.metric("Class Rank", f"{rank} / {total}")
                else:
                    st.metric("Class Rank", "Not Available")
            
            # Student insights
            st.subheader(f"💡 Performance Insights for {selected_student}")
            insights = get_student_performance_insights(student_data, subjects, data['subject_stats'])
            for insight in insights:
                st.write(insight)
            
            # Student charts
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Subject Performance Radar**")
                radar_img = create_student_performance_chart(student_data, subjects)
                st.image(radar_img)
            
            with col2:
                st.write("**Subject vs Class Average**")
                bar_img = create_student_subject_bar_chart(student_data, subjects, data['subject_stats'])
                st.image(bar_img)
            
            # Detailed subject scores
            st.subheader("📋 Detailed Subject Scores")
            subject_scores_data = []
            for subject in subjects:
                pct_col = subject + "_pct"
                student_score = student_data.get(pct_col, np.nan)
                class_avg = data['subject_stats'].get(subject, {}).get('average', np.nan)
                
                if not pd.isna(student_score) and not pd.isna(class_avg):
                    difference = student_score - class_avg
                    status = "Above Average" if difference > 0 else "Below Average" if difference < 0 else "At Average"
                    subject_scores_data.append({
                        'Subject': subject,
                        'Student Score': f"{student_score:.2f}%",
                        'Class Average': f"{class_avg:.2f}%",
                        'Difference': f"{difference:+.2f}%",
                        'Status': status
                    })
            
            if subject_scores_data:
                subject_scores_df = pd.DataFrame(subject_scores_data)
                st.dataframe(subject_scores_df, use_container_width=True)
    else:
        st.warning("Student Name column not found for individual analysis.")

    # Data Preview
    with st.expander("🔍 View Processed Data Preview"):
        preview_cols = ["Student Name"] + subjects + [s + "_pct" for s in subjects] + ["Overall_pct"]
        available_cols = [c for c in preview_cols if c in processed_df.columns]
        st.dataframe(processed_df[available_cols].head(15), use_container_width=True)

    with st.expander("📊 View Detailed Statistics"):
        st.write("**Subject-wise Statistics:**")
        stats_data = []
        for subject in subjects:
            pct_col = subject + "_pct"
            if pct_col in processed_df.columns:
                subject_data = processed_df[pct_col].dropna()
                if len(subject_data) > 0:
                    stats_data.append({
                        "Subject": subject,
                        "Average": subject_data.mean(),
                        "Median": subject_data.median(),
                        "Highest": subject_data.max(),
                        "Lowest": subject_data.min(),
                        "Std Dev": subject_data.std(),
                        "Count": len(subject_data)
                    })
        stats_df = pd.DataFrame(stats_data)
        if not stats_df.empty:
            # Format numeric columns
            numeric_cols = ['Average', 'Median', 'Highest', 'Lowest', 'Std Dev']
            for col in numeric_cols:
                if col in stats_df.columns:
                    stats_df[col] = stats_df[col].round(2)
            st.dataframe(stats_df, use_container_width=True)
        else:
            st.info("No subject statistics available.")

    # PDF Download Button for Single File
    st.markdown("---")
    st.subheader("📄 Download Report")
    
    output_type = st.radio("Choose output format:", 
                          ("PDF Report (Recommended)", "Editable Word Document (.docx)"),
                          horizontal=True)

    if st.button("📥 Generate & Download Report", type="primary", key="single_download"):
        logo_bytes = None
        if logo_file is not None:
            logo_bytes = logo_file.getvalue() if hasattr(logo_file, 'getvalue') else logo_file

        class_name = processed_df.get("Class Name", "Unknown Class").iloc[0] if "Class Name" in processed_df.columns else "Unknown_Class"
        class_name_safe = re.sub(r'[^\w\-_]', '', str(class_name).replace(" ", "_"))
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        try:
            if output_type == "PDF Report (Recommended)":
                pdf_buffer = generate_pdf_buffer(processed_df, subjects, logo_bytes)
                filename = f"student_report_{class_name_safe}_{timestamp}.pdf"
                
                st.download_button(
                    label="📥 Download PDF Report",
                    data=pdf_buffer,
                    file_name=filename,
                    mime="application/pdf",
                    type="primary"
                )
            else:  # DOCX
                docx_buffer = generate_docx_buffer(processed_df, subjects, logo_bytes)
                filename = f"student_report_{class_name_safe}_{timestamp}.docx"
                
                st.download_button(
                    label="📥 Download Word Document",
                    data=docx_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
            
            st.success("✅ Report generated successfully! Click the download button above.")
            
        except Exception as e:
            st.error(f"❌ Error generating report: {e}")

else:
    # Two-file comparison mode (existing code remains the same)
    # ... [rest of the comparison mode code remains unchanged]
     # Two-file comparison mode
    st.header("📈 Two-Week Comparison Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📅 Week 1 Data")
        week1_file = st.file_uploader("Upload Week 1 Excel file", type=["xlsx"], key="week1")
        week1_name = st.text_input("Week 1 Name", value="Week 1")
        
    with col2:
        st.subheader("📅 Week 2 Data")
        week2_file = st.file_uploader("Upload Week 2 Excel file", type=["xlsx"], key="week2")
        week2_name = st.text_input("Week 2 Name", value="Week 2")
    
    if week1_file and week2_file:
        try:
            # Load and process both files
            with st.spinner("Loading and processing both files..."):
                df_week1 = pd.read_excel(week1_file)
                df_week2 = pd.read_excel(week2_file)
                
                # Detect subjects for both files
                ignore_cols = IGNORE_COLS_DEFAULT.copy()
                subjects_week1 = [c for c in df_week1.columns if c not in ignore_cols and c.lower() not in [ic.lower() for ic in ignore_cols]]
                subjects_week2 = [c for c in df_week2.columns if c not in ignore_cols and c.lower() not in [ic.lower() for ic in ignore_cols]]
                
                # Use common subjects
                common_subjects = list(set(subjects_week1) & set(subjects_week2))
                if not common_subjects:
                    st.error("❌ No common subjects found between the two files.")
                    st.stop()
                
                st.success(f"✅ Found {len(common_subjects)} common subjects: {common_subjects}")
                
                # Process both datasets
                week1_data = process_dataframe(df_week1, common_subjects)
                week2_data = process_dataframe(df_week2, common_subjects)
                
                # Perform comparison
                comparison = compare_weeks(week1_data, week2_data, week1_name, week2_name)
                st.session_state.comparison_data = comparison
                
            # Display comparison results
            st.markdown("---")
            st.header("📊 Comparison Results")
            
            # Key metrics comparison
            st.subheader("📈 Key Metrics Comparison")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                delta_class_avg = comparison['class_avg_change']
                st.metric(
                    "Class Average",
                    f"{week2_data['class_avg']:.2f}%",
                    f"{delta_class_avg:+.2f}%",
                    delta_color="normal" if delta_class_avg >= 0 else "inverse"
                )
            
            with col2:
                delta_highest = comparison['highest_change']
                st.metric(
                    "Highest Score",
                    f"{week2_data['highest_pct']:.2f}%",
                    f"{delta_highest:+.2f}%",
                    delta_color="normal" if delta_highest >= 0 else "inverse"
                )
            
            with col3:
                delta_strong = comparison['strong_students_change']
                st.metric(
                    "Strong Students",
                    len(week2_data['strong_students']),
                    f"{delta_strong:+d}",
                    delta_color="normal" if delta_strong >= 0 else "inverse"
                )
            
            with col4:
                delta_weak = comparison['weak_students_change']
                st.metric(
                    "Weak Students",
                    len(week2_data['weak_students']),
                    f"{delta_weak:+d}",
                    delta_color="normal" if delta_weak <= 0 else "inverse"
                )
            
            # Comparison charts
            st.subheader("📊 Comparison Charts")
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Overall Metrics Comparison**")
                comp_chart = create_comparison_chart(week1_data, week2_data, week1_name, week2_name)
                st.image(comp_chart)
            
            with col2:
                st.write("**Performance Category Changes**")
                cat_data = {
                    'Strong': [len(week1_data['strong_students']), len(week2_data['strong_students'])],
                    'Average': [len(week1_data['avg_students']), len(week2_data['avg_students'])],
                    'Weak': [len(week1_data['weak_students']), len(week2_data['weak_students'])]
                }
                
                fig, ax = plt.subplots(figsize=(8, 6))
                x = np.arange(3)
                width = 0.35
                
                bars1 = ax.bar(x - width/2, [cat_data[cat][0] for cat in ['Strong', 'Average', 'Weak']], 
                              width, label=week1_name, color='#3498db', alpha=0.7)
                bars2 = ax.bar(x + width/2, [cat_data[cat][1] for cat in ['Strong', 'Average', 'Weak']], 
                              width, label=week2_name, color='#2ecc71', alpha=0.7)
                
                ax.set_ylabel('Number of Students')
                ax.set_title('Performance Category Changes')
                ax.set_xticks(x)
                ax.set_xticklabels(['Strong', 'Average', 'Weak'])
                ax.legend()
                ax.grid(axis='y', alpha=0.3)
                
                for bars in [bars1, bars2]:
                    for bar in bars:
                        height = bar.get_height()
                        ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                               f'{int(height)}', ha='center', va='bottom', fontweight='bold')
                
                st.pyplot(fig)
            
            # Subject-wise comparison
            st.subheader("📚 Subject-wise Performance Changes")
            subject_comp_chart = create_subject_comparison_chart(comparison['subject_comparison'], week1_name, week2_name)
            st.image(subject_comp_chart)
            
            # Detailed subject comparison table
            st.subheader("📋 Subject-wise Detailed Comparison")
            subject_comparison_data = []
            for subject, stats in comparison['subject_comparison'].items():
                subject_comparison_data.append({
                    'Subject': subject,
                    f'{week1_name} Avg': f"{stats['week1_avg']:.2f}%",
                    f'{week2_name} Avg': f"{stats['week2_avg']:.2f}%",
                    'Change': f"{stats['change']:+.2f}%",
                    'Change %': f"{stats['change_pct']:+.1f}%"
                })
            
            subject_df = pd.DataFrame(subject_comparison_data)
            st.dataframe(subject_df, use_container_width=True)
            
            # Individual Student Comparison Section
            if comparison['common_students_count'] > 0:
                st.markdown("---")
                st.header("👥 Individual Student Comparison")
                
                # Student selector
                student_names = comparison['common_students']['name'].tolist()
                selected_student = st.selectbox("Select Student for Detailed Analysis", student_names)
                
                if selected_student:
                    student_data = comparison['student_subject_comparison'][selected_student]
                    student_overall = comparison['common_students'][comparison['common_students']['name'] == selected_student].iloc[0]
                    
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.subheader(f"📊 {selected_student} - Subject Performance")
                        student_chart = create_student_comparison_chart(student_data, week1_name, week2_name)
                        st.image(student_chart)
                    
                    with col2:
                        st.subheader("Overall Performance")
                        st.metric(
                            "Overall Score",
                            f"{student_overall['week2_score']:.2f}%",
                            f"{student_overall['change']:+.2f}%",
                            delta_color="normal" if student_overall['change'] >= 0 else "inverse"
                        )
                        
                        st.write("**Subject Changes:**")
                        for subject, scores in student_data.items():
                            change_icon = "📈" if scores['change'] >= 0 else "📉"
                            st.write(f"{change_icon} {subject}: {scores['change']:+.1f}%")
                
                # Top and Bottom Improvers
                st.subheader("🏆 Student Progress Leaders")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Top 5 Improvers**")
                    top_improvers = comparison['common_students'].nlargest(5, 'change')
                    display_top = top_improvers[['name', 'week1_score', 'week2_score', 'change']].copy()
                    display_top.columns = ['Student', f'{week1_name}', f'{week2_name}', 'Change']
                    display_top[f'{week1_name}'] = display_top[f'{week1_name}'].round(2)
                    display_top[f'{week2_name}'] = display_top[f'{week2_name}'].round(2)
                    display_top['Change'] = display_top['Change'].round(2)
                    st.dataframe(display_top, use_container_width=True)
                
                with col2:
                    st.write("**Top 5 Decliners**")
                    bottom_improvers = comparison['common_students'].nsmallest(5, 'change')
                    display_bottom = bottom_improvers[['name', 'week1_score', 'week2_score', 'change']].copy()
                    display_bottom.columns = ['Student', f'{week1_name}', f'{week2_name}', 'Change']
                    display_bottom[f'{week1_name}'] = display_bottom[f'{week1_name}'].round(2)
                    display_bottom[f'{week2_name}'] = display_bottom[f'{week2_name}'].round(2)
                    display_bottom['Change'] = display_bottom['Change'].round(2)
                    st.dataframe(display_bottom, use_container_width=True)
                
                # Student progress distribution
                st.write("**Student Progress Distribution**")
                fig, ax = plt.subplots(figsize=(10, 6))
                changes = comparison['common_students']['change']
                n, bins, patches = ax.hist(changes, bins=15, color='lightblue', alpha=0.7, edgecolor='black')
                ax.set_xlabel('Score Change (%)')
                ax.set_ylabel('Number of Students')
                ax.set_title('Distribution of Student Score Changes')
                ax.grid(axis='y', alpha=0.3)
                ax.axvline(x=0, color='red', linestyle='--', alpha=0.7, label='No Change')
                ax.legend()
                
                for i, (count, patch) in enumerate(zip(n, patches)):
                    if count > 0:
                        ax.text(patch.get_x() + patch.get_width()/2, count + 0.1,
                               f'{int(count)}', ha='center', va='bottom', fontweight='bold')
                
                st.pyplot(fig)
            
            # Summary insights
            st.markdown("---")
            st.subheader("💡 Summary Insights")
            
            insights = []
            if comparison['class_avg_change'] > 0:
                insights.append(f"✅ **Overall Improvement**: Class average increased by {comparison['class_avg_change']:.2f}% ({comparison['class_avg_change_pct']:.1f}%)")
            else:
                insights.append(f"⚠️ **Overall Decline**: Class average decreased by {abs(comparison['class_avg_change']):.2f}% ({abs(comparison['class_avg_change_pct']):.1f}%)")
            
            if comparison['strong_students_change'] > 0:
                insights.append(f"✅ **More Strong Performers**: Number of strong students increased by {comparison['strong_students_change']}")
            else:
                insights.append(f"⚠️ **Fewer Strong Performers**: Number of strong students decreased by {abs(comparison['strong_students_change'])}")
            
            # Find best and worst performing subjects
            if comparison['subject_comparison']:
                best_subject = max(comparison['subject_comparison'].items(), key=lambda x: x[1]['change'])
                worst_subject = min(comparison['subject_comparison'].items(), key=lambda x: x[1]['change'])
                
                insights.append(f"🎯 **Best Improvement**: {best_subject[0]} improved by {best_subject[1]['change']:+.2f}%")
                insights.append(f"📉 **Largest Decline**: {worst_subject[0]} declined by {abs(worst_subject[1]['change']):.2f}%")
            
            for insight in insights:
                st.write(insight)
            
            # PDF Download Button
            st.markdown("---")
            st.subheader("📄 Download Comparison Report")
            
            if st.button("📥 Generate & Download Comparison PDF Report", type="primary"):
                logo_bytes = None
                if logo_file is not None:
                    logo_bytes = logo_file.getvalue() if hasattr(logo_file, 'getvalue') else logo_file
                
                with st.spinner("Generating comprehensive PDF report..."):
                    try:
                        pdf_buffer = generate_comparison_pdf_buffer(
                            week1_data, week2_data, comparison, week1_name, week2_name, logo_bytes
                        )
                        
                        # Create safe filename
                        class_name = week1_data['df'].get('Class Name', 'Unknown Class').iloc[0] if 'Class Name' in week1_data['df'].columns else 'Unknown_Class'
                        class_name_safe = re.sub(r'[^\w\-_]', '', str(class_name).replace(" ", "_"))
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        filename = f"comparison_report_{class_name_safe}_{timestamp}.pdf"
                        
                        st.download_button(
                            label="📥 Download PDF Report",
                            data=pdf_buffer,
                            file_name=filename,
                            mime="application/pdf",
                            type="primary"
                        )
                        
                        st.success("✅ PDF report generated successfully! Click the download button above.")
                        
                    except Exception as e:
                        st.error(f"❌ Error generating PDF report: {e}")
                
        except Exception as e:
            st.error(f"❌ Error processing comparison: {e}")
            st.stop()
    else:
        st.info("👆 Please upload both Week 1 and Week 2 Excel files to begin comparison analysis.")

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: gray;'>"
    "Student Performance Analyzer • Built with Streamlit • "
    f"Report generated on {report_date}"
    "</div>",
    unsafe_allow_html=True
)